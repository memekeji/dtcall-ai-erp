from django.shortcuts import render, get_object_or_404, redirect
from django.urls import reverse
from django.contrib.auth.decorators import login_required
import os
from django.conf import settings
from django.contrib import messages
from django.core.paginator import Paginator
from django.db.models import Q
from django.http import JsonResponse, HttpResponse
from django.utils import timezone
from django.views.decorators.http import require_http_methods
import json
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
import tempfile

# 导入网盘中的Office文档预览处理器
from apps.disk.utils.office_preview import OfficePreviewHandler

from .models import (
    PersonalSchedule, WorkRecord, WorkReport, 
    PersonalNote, PersonalTask, PersonalContact, MeetingMinutes
)
from apps.oa.constants import MeetingTypeChoices
from .forms import (
    PersonalScheduleForm, WorkRecordForm, WorkReportForm,
    PersonalNoteForm, PersonalTaskForm, PersonalContactForm, MeetingMinutesForm
)


@login_required
def schedule_list(request):
    """日程安排列表"""
    date = request.GET.get('date', '')
    if not date:
        date = timezone.now().date()
    else:
        date = datetime.strptime(date, '%Y-%m-%d').date()
    
    schedules = PersonalSchedule.objects.filter(
        user=request.user,
        start_time__date=date
    ).order_by('start_time')
    
    context = {
        'schedules': schedules,
        'current_date': date,
    }
    return render(request, 'personal/schedule/list.html', context)


@login_required
def schedule_calendar(request):
    """日程日历"""
    year = int(request.GET.get('year', timezone.now().year))
    month = int(request.GET.get('month', timezone.now().month))
    
    # 获取当月的所有日程
    start_date = datetime(year, month, 1).date()
    if month == 12:
        end_date = datetime(year + 1, 1, 1).date()
    else:
        end_date = datetime(year, month + 1, 1).date()
    
    schedules = PersonalSchedule.objects.filter(
        user=request.user,
        start_time__date__gte=start_date,
        start_time__date__lt=end_date
    ).order_by('start_time')
    
    # 按日期分组
    schedule_dict = {}
    for schedule in schedules:
        date_key = schedule.start_time.date().strftime('%Y-%m-%d')
        if date_key not in schedule_dict:
            schedule_dict[date_key] = []
        schedule_dict[date_key].append(schedule)
    
    context = {
        'year': year,
        'month': month,
        'schedule_dict': schedule_dict,
        'schedules': schedules,
    }
    return render(request, 'personal/schedule/calendar.html', context)


@login_required
def schedule_form(request, pk=None):
    """日程表单"""
    schedule = None
    if pk:
        schedule = get_object_or_404(PersonalSchedule, pk=pk, user=request.user)
    
    if request.method == 'POST':
        form = PersonalScheduleForm(request.POST, instance=schedule)
        if form.is_valid():
            try:
                schedule = form.save(commit=False)
                schedule.user = request.user
                schedule.save()
                messages.success(request, '日程保存成功！')
                return redirect('personal:schedule_list')
            except Exception as e:
                messages.error(request, f'保存日程时出错：{str(e)}')
        else:
            # 显示表单错误信息
            for field, errors in form.errors.items():
                for error in errors:
                    messages.error(request, f'{field}: {error}')
    else:
        # 如果是新建日程，设置默认值
        if not schedule:
            now = timezone.now()
            initial = {
                'start_time': now.strftime('%Y-%m-%dT%H:%M'),
                'end_time': (now + timedelta(hours=1)).strftime('%Y-%m-%dT%H:%M'),
                'status': 'pending',
                'priority': 2
            }
            form = PersonalScheduleForm(initial=initial)
        else:
            form = PersonalScheduleForm(instance=schedule)
    
    context = {'form': form, 'schedule': schedule}
    return render(request, 'personal/schedule/form.html', context)


@login_required
def work_record_list(request):
    """工作记录列表"""
    search = request.GET.get('search', '')
    work_type = request.GET.get('work_type', '')
    date_from = request.GET.get('date_from', '')
    date_to = request.GET.get('date_to', '')
    
    records = WorkRecord.objects.filter(user=request.user)
    
    if search:
        records = records.filter(
            Q(title__icontains=search) | 
            Q(content__icontains=search)
        )
    
    if work_type:
        records = records.filter(work_type=work_type)
    
    if date_from:
        records = records.filter(work_date__gte=date_from)
    
    if date_to:
        records = records.filter(work_date__lte=date_to)
    
    records = records.order_by('-work_date', '-start_time')
    
    # 分页
    paginator = Paginator(records, 20)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    context = {
        'page_obj': page_obj,
        'search': search,
        'work_type': work_type,
        'date_from': date_from,
        'date_to': date_to,
        'work_types': WorkRecord.WORK_TYPES,
    }
    return render(request, 'personal/record/list.html', context)


@login_required
def work_record_form(request, pk=None):
    """工作记录表单"""
    record = None
    if pk:
        record = get_object_or_404(WorkRecord, pk=pk, user=request.user)
    
    if request.method == 'POST':
        form = WorkRecordForm(request.POST, instance=record)
        if form.is_valid():
            record = form.save(commit=False)
            record.user = request.user
            record.department = request.user.department
            record.save()
            messages.success(request, '工作记录保存成功！')
            return redirect('personal:work_record_list')
    else:
        form = WorkRecordForm(instance=record)
    
    context = {'form': form, 'record': record}
    return render(request, 'personal/record/form.html', context)


@login_required
def work_report_list(request):
    """工作汇报列表"""
    search = request.GET.get('search', '')
    report_type = request.GET.get('report_type', '')
    
    reports = WorkReport.objects.filter(user=request.user)
    
    if search:
        reports = reports.filter(title__icontains=search)
    
    if report_type:
        reports = reports.filter(report_type=report_type)
    
    reports = reports.order_by('-report_date')
    
    # 分页
    paginator = Paginator(reports, 20)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    context = {
        'page_obj': page_obj,
        'search': search,
        'report_type': report_type,
        'report_types': WorkReport.REPORT_TYPES,
    }
    return render(request, 'personal/report/list.html', context)


@login_required
def work_report_form(request, pk=None):
    """工作汇报表单"""
    report = None
    if pk:
        report = get_object_or_404(WorkReport, pk=pk, user=request.user)
    
    if request.method == 'POST':
        form = WorkReportForm(request.POST, instance=report)
        if form.is_valid():
            report = form.save(commit=False)
            report.user = request.user
            report.department = request.user.department
            
            # 处理提交状态
            if 'submit' in request.POST:
                report.is_submitted = True
                report.submitted_at = timezone.now()
            
            report.save()
            
            # 处理接收人
            recipient_ids = request.POST.getlist('recipient_users')
            report.recipient_users.set(recipient_ids)
            
            messages.success(request, '工作汇报保存成功！')
            return redirect('personal:work_report_list')
    else:
        form = WorkReportForm(instance=report)
    
    context = {'form': form, 'report': report}
    return render(request, 'personal/report/form.html', context)


@login_required
def note_list(request):
    """个人笔记列表"""
    search = request.GET.get('search', '')
    category = request.GET.get('category', '')
    
    notes = PersonalNote.objects.filter(user=request.user)
    
    if search:
        notes = notes.filter(
            Q(title__icontains=search) | 
            Q(content__icontains=search) |
            Q(tags__icontains=search)
        )
    
    if category:
        notes = notes.filter(category=category)
    
    notes = notes.order_by('-is_important', '-updated_at')
    
    # 分页
    paginator = Paginator(notes, 20)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    context = {
        'page_obj': page_obj,
        'search': search,
        'category': category,
        'categories': PersonalNote.CATEGORIES,
    }
    return render(request, 'personal/note/list.html', context)


@login_required
def note_form(request, pk=None):
    """个人笔记表单"""
    note = None
    if pk:
        note = get_object_or_404(PersonalNote, pk=pk, user=request.user)
    
    if request.method == 'POST':
        form = PersonalNoteForm(request.POST, instance=note)
        if form.is_valid():
            note = form.save(commit=False)
            note.user = request.user
            note.save()
            messages.success(request, '笔记保存成功！')
            return redirect('personal:note_list')
    else:
        form = PersonalNoteForm(instance=note)
    
    context = {'form': form, 'note': note}
    return render(request, 'personal/note/form.html', context)


@login_required
def task_list(request):
    """个人任务列表"""
    status = request.GET.get('status', '')
    priority = request.GET.get('priority', '')
    
    tasks = PersonalTask.objects.filter(user=request.user)
    
    if status:
        tasks = tasks.filter(status=status)
    
    if priority:
        tasks = tasks.filter(priority=priority)
    
    tasks = tasks.order_by('-priority', 'due_date')
    
    # 分页
    paginator = Paginator(tasks, 20)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    context = {
        'page_obj': page_obj,
        'status': status,
        'priority': priority,
        'status_choices': PersonalTask.STATUS_CHOICES,
        'priority_choices': PersonalTask.PRIORITY_CHOICES,
    }
    return render(request, 'personal/task/list.html', context)


@login_required
def task_form(request, pk=None):
    """个人任务表单"""
    task = None
    if pk:
        task = get_object_or_404(PersonalTask, pk=pk, user=request.user)
    
    if request.method == 'POST':
        form = PersonalTaskForm(request.POST, instance=task)
        if form.is_valid():
            task = form.save(commit=False)
            task.user = request.user
            
            # 如果状态改为完成，设置完成时间
            if task.status == 'completed' and not task.completed_at:
                task.completed_at = timezone.now()
            
            task.save()
            messages.success(request, '任务保存成功！')
            return redirect('personal:task_list')
    else:
        form = PersonalTaskForm(instance=task)
    
    context = {'form': form, 'task': task}
    return render(request, 'personal/task/form.html', context)


@login_required
def contact_list(request):
    """个人通讯录列表"""
    search = request.GET.get('search', '')
    
    contacts = PersonalContact.objects.filter(user=request.user)
    
    if search:
        contacts = contacts.filter(
            Q(name__icontains=search) |
            Q(company__icontains=search) |
            Q(phone__icontains=search) |
            Q(mobile__icontains=search) |
            Q(email__icontains=search)
        )
    
    contacts = contacts.order_by('-is_important', 'name')
    
    # 分页
    paginator = Paginator(contacts, 20)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    context = {
        'page_obj': page_obj,
        'search': search,
    }
    return render(request, 'personal/contact/list.html', context)


@login_required
def contact_form(request, pk=None):
    """个人通讯录表单"""
    contact = None
    if pk:
        contact = get_object_or_404(PersonalContact, pk=pk, user=request.user)
    
    if request.method == 'POST':
        form = PersonalContactForm(request.POST, instance=contact)
        if form.is_valid():
            contact = form.save(commit=False)
            contact.user = request.user
            contact.save()
            messages.success(request, '联系人保存成功！')
            return redirect('personal:contact_list')
    else:
        form = PersonalContactForm(instance=contact)
    
    context = {'form': form, 'contact': contact}
    return render(request, 'personal/contact/form.html', context)


@login_required
@require_http_methods(["POST"])
def task_toggle_status(request, pk):
    """切换任务状态"""
    task = get_object_or_404(PersonalTask, pk=pk, user=request.user)
    
    if task.status == 'completed':
        task.status = 'todo'
        task.completed_at = None
        message = '任务已标记为待办'
    else:
        task.status = 'completed'
        task.completed_at = timezone.now()
        task.progress = 100
        message = '任务已标记为完成'
    
    task.save()
    
    return JsonResponse({
        'success': True,
        'message': message,
        'status': task.status
    })


@login_required
def dashboard(request):
    """个人办公首页"""
    today = timezone.now().date()
    
    # 今日日程
    today_schedules = PersonalSchedule.objects.filter(
        user=request.user,
        start_time__date=today
    ).order_by('start_time')[:5]
    
    # 待办任务
    pending_tasks = PersonalTask.objects.filter(
        user=request.user,
        status__in=['todo', 'in_progress']
    ).order_by('-priority', 'due_date')[:5]
    
    # 最近笔记
    recent_notes = PersonalNote.objects.filter(
        user=request.user
    ).order_by('-updated_at')[:5]
    
    # 统计数据
    stats = {
        'total_tasks': PersonalTask.objects.filter(user=request.user).count(),
        'completed_tasks': PersonalTask.objects.filter(user=request.user, status='completed').count(),
        'pending_tasks': PersonalTask.objects.filter(user=request.user, status__in=['todo', 'in_progress']).count(),
        'total_notes': PersonalNote.objects.filter(user=request.user).count(),
    }
    
    context = {
        'today_schedules': today_schedules,
        'pending_tasks': pending_tasks,
        'recent_notes': recent_notes,
        'stats': stats,
    }
    return render(request, 'personal/dashboard.html', context)


@login_required
def schedule_delete(request, pk):
    """删除日程"""
    schedule = get_object_or_404(PersonalSchedule, pk=pk, user=request.user)
    if request.method == 'POST':
        schedule.delete()
        messages.success(request, '日程删除成功！')
        return redirect('personal:schedule_list')
    return render(request, 'personal/schedule/delete.html', {'schedule': schedule})


@login_required
def work_record_delete(request, pk):
    """删除工作记录"""
    record = get_object_or_404(WorkRecord, pk=pk, user=request.user)
    if request.method == 'POST':
        record.delete()
        messages.success(request, '工作记录删除成功！')
        return redirect('personal:work_record_list')
    return render(request, 'personal/record/delete.html', {'record': record})


@login_required
def work_calendar(request):
    """工作日历"""
    year = int(request.GET.get('year', timezone.now().year))
    month = int(request.GET.get('month', timezone.now().month))
    
    # 获取当月的所有工作记录
    start_date = datetime(year, month, 1).date()
    if month == 12:
        end_date = datetime(year + 1, 1, 1).date()
    else:
        end_date = datetime(year, month + 1, 1).date()
    
    records = WorkRecord.objects.filter(
        user=request.user,
        work_date__gte=start_date,
        work_date__lt=end_date
    ).order_by('work_date')
    
    context = {
        'year': year,
        'month': month,
        'records': records,
    }
    return render(request, 'personal/record/calendar.html', context)


@login_required
def work_report_delete(request, pk):
    """删除工作汇报"""
    report = get_object_or_404(WorkReport, pk=pk, user=request.user)
    if request.method == 'POST':
        report.delete()
        messages.success(request, '工作汇报删除成功！')
        return redirect('personal:work_report_list')
    return render(request, 'personal/report/delete.html', {'report': report})


@login_required
def work_report_detail(request, pk):
    """工作汇报详情"""
    report = get_object_or_404(WorkReport, pk=pk, user=request.user)
    context = {'report': report}
    return render(request, 'personal/report/detail.html', context)


@login_required
def note_delete(request, pk):
    """删除笔记"""
    note = get_object_or_404(PersonalNote, pk=pk, user=request.user)
    if request.method == 'POST':
        note.delete()
        messages.success(request, '笔记删除成功！')
        return redirect('personal:note_list')
    return render(request, 'personal/note/delete.html', {'note': note})


@login_required
def task_delete(request, pk):
    """删除任务"""
    task = get_object_or_404(PersonalTask, pk=pk, user=request.user)
    if request.method == 'POST':
        task.delete()
        messages.success(request, '任务删除成功！')
        return redirect('personal:task_list')
    return render(request, 'personal/task/delete.html', {'task': task})


@login_required
def contact_delete(request, pk):
    """删除联系人"""
    contact = get_object_or_404(PersonalContact, pk=pk, user=request.user)
    if request.method == 'POST':
        contact.delete()
        messages.success(request, '联系人删除成功！')
        return redirect('personal:contact_list')
    return render(request, 'personal/contact/delete.html', {'contact': contact})


@login_required
def notice_list(request):
    """公告通知列表"""
    from apps.system.models import Notice
    from django.core.paginator import Paginator
    
    search = request.GET.get('search', '')
    notice_type = request.GET.get('notice_type', '')
    
    # 获取已发布的公告
    notices = Notice.objects.filter(is_published=True).select_related('author')
    
    if search:
        notices = notices.filter(
            Q(title__icontains=search) | Q(content__icontains=search)
        )
    
    if notice_type:
        notices = notices.filter(notice_type=notice_type)
    
    notices = notices.order_by('-is_top', '-publish_time')
    
    # 分页
    paginator = Paginator(notices, 10)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    context = {
        'page_obj': page_obj,
        'search': search,
        'notice_type': notice_type,
    }
    return render(request, 'personal/notice/list.html', context)


@login_required
def news_list(request):
    """个人办公 - 公司新闻列表（只读）"""
    from apps.system.models import Notice
    from django.core.paginator import Paginator
    
    search = request.GET.get('search', '')
    
    # 获取已发布的公司新闻
    notices = Notice.objects.filter(
        notice_type='company',
        is_published=True
    ).select_related('author')
    
    if search:
        notices = notices.filter(
            Q(title__icontains=search) | Q(content__icontains=search)
        )
    
    notices = notices.order_by('-is_top', '-publish_time')
    
    # 分页
    paginator = Paginator(notices, 10)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    context = {
        'page_obj': page_obj,
        'search': search,
    }
    return render(request, 'personal/news/list.html', context)


@login_required
def minutes_list(request):
    """会议纪要列表"""
    search = request.GET.get('search', '')
    meeting_type = request.GET.get('meeting_type', '')
    date_from = request.GET.get('date_from', '')
    date_to = request.GET.get('date_to', '')
    
    minutes = MeetingMinutes.objects.filter(Q(user=request.user) | Q(is_public=True))
    
    if search:
        minutes = minutes.filter(
            Q(title__icontains=search) | 
            Q(content__icontains=search) |
            Q(decisions__icontains=search)
        )
    
    if meeting_type:
        minutes = minutes.filter(meeting_type=meeting_type)
    
    if date_from:
        minutes = minutes.filter(meeting_date__date__gte=date_from)
    
    if date_to:
        minutes = minutes.filter(meeting_date__date__lte=date_to)
    
    minutes = minutes.order_by('-meeting_date')
    
    # 分页
    paginator = Paginator(minutes, 20)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    context = {
        'page_obj': page_obj,
        'search': search,
        'meeting_type': meeting_type,
        'date_from': date_from,
        'date_to': date_to,
        'meeting_types': MeetingTypeChoices.choices,
    }
    return render(request, 'personal/minutes/list.html', context)


@login_required
def minutes_form(request, pk=None):
    """会议纪要表单处理"""
    # 检查是否有预约ID参数
    reservation_id = request.GET.get('reservation_id')
    
    # 获取或创建会议纪要对象
    minute = get_object_or_404(MeetingMinutes, pk=pk) if pk else None
    
    # 检查权限：只有创建者可以编辑
    if minute and minute.user != request.user:
        messages.error(request, '您没有权限编辑此会议纪要！')
        return redirect('personal:minutes_list')
    
    if request.method == 'POST':
        # 处理AI生成纪要请求
        if request.POST.get('generate_ai_minutes') == 'true':
            try:
                # 获取会议ID和预约ID
                meeting_id = request.POST.get('meeting_id')
                reservation_id = request.POST.get('reservation_id') or request.GET.get('reservation_id')
                
                # 如果会议ID为空但有预约ID，尝试基于预约ID创建或获取会议记录
                if not meeting_id and reservation_id:
                    try:
                        from apps.system.models import MeetingReservation
                        from apps.oa.models import MeetingRecord
                        
                        # 获取预约记录
                        reservation = MeetingReservation.objects.get(pk=reservation_id)
                        
                        # 尝试查找关联的会议记录
                        meeting_record = MeetingRecord.objects.filter(
                            title=reservation.title,
                            meeting_date=reservation.start_time
                        ).first()
                        
                        # 如果没有找到，创建新的会议记录
                        if not meeting_record:
                            meeting_room_name = reservation.meeting_room.name if reservation.meeting_room else ''
                            meeting_record = MeetingRecord.objects.create(
                                title=reservation.title,
                                meeting_date=reservation.start_time,
                                location=meeting_room_name,
                                host=reservation.organizer,
                                content=reservation.description or ''
                            )
                            # 添加参与者 - 确保正确处理ManyToMany关系
                            participants_list = list(reservation.attendees.all())
                            if participants_list:
                                meeting_record.participants.set(participants_list)
                            
                            import logging
                            logger = logging.getLogger(__name__)
                            logger.info(f"AI生成时基于预约ID {reservation_id} 自动创建了会议记录 {meeting_record.id}")
                        
                        # 使用自动创建或找到的会议记录ID
                        meeting_id = meeting_record.id
                        import logging
                        logger = logging.getLogger(__name__)
                        logger.info(f"从预约ID {reservation_id} 关联到会议记录ID {meeting_id}")
                    except Exception as e:
                        import logging
                        logger = logging.getLogger(__name__)
                        logger.error(f"基于预约ID创建或获取会议记录失败: {str(e)}")
                        return JsonResponse({'code': 1, 'msg': '无法基于预约信息创建会议记录，请手动选择会议记录'})
                
                # 如果仍然没有会议ID，返回错误
                if not meeting_id:
                    return JsonResponse({'code': 1, 'msg': '会议ID不能为空，请先选择一个会议记录'})
                
                # 导入MeetingRecord模型
                from apps.oa.models import MeetingRecord
                
                # 尝试获取会议记录
                try:
                    meeting_record = MeetingRecord.objects.get(pk=meeting_id)
                    # 安全地检查用户权限：主持人、记录员或参与者
                    if not (meeting_record.host == request.user or 
                            meeting_record.recorder == request.user or 
                            meeting_record.participants.filter(id=request.user.id).exists()):
                        return JsonResponse({'code': 1, 'msg': '您没有权限为此会议生成纪要'})
                except MeetingRecord.DoesNotExist:
                    import logging
                    logger = logging.getLogger(__name__)
                    logger.error(f"会议记录不存在：会议ID {meeting_id} 未找到对应的MeetingRecord")
                    
                    # 检查是否可能是reservation_id，如果是，给出更友好的错误提示
                    if reservation_id and meeting_id == reservation_id:
                        return JsonResponse({'code': 1, 'msg': '未找到对应的会议记录。请注意：会议室预约ID和会议记录ID是不同的概念，请先创建会议记录，或者从会议记录页面进入纪要编辑'})
                    else:
                        # 检查是否有对应的会议室预约，提供更详细的提示
                        try:
                            from apps.system.models import MeetingReservation
                            reservation = MeetingReservation.objects.filter(pk=meeting_id).first()
                            if reservation:
                                return JsonResponse({'code': 1, 'msg': '您提供的ID似乎是会议室预约ID，而不是会议记录ID。请先基于此预约创建会议记录，然后再生成纪要'})
                        except Exception:
                            pass
                        
                        return JsonResponse({'code': 1, 'msg': f'会议记录不存在：会议ID {meeting_id} 未找到，请确认您已选择了正确的会议记录'})
                
                # 处理音频文件上传
                audio_file = request.FILES.get('audio_file')
                audio_file_path = None
                
                if audio_file:
                    # 检查文件类型
                    import os
                    ext = os.path.splitext(audio_file.name)[1].lower()
                    allowed_extensions = ['.mp3', '.wav', '.flac', '.ogg', '.aac']
                    if ext not in allowed_extensions:
                        return JsonResponse({'code': 1, 'msg': '不支持的音频格式，请上传mp3、wav、flac、ogg或aac格式'})
                    
                    # 生成唯一文件名
                    timestamp = timezone.now().strftime('%Y%m%d%H%M%S')
                    filename = f"meeting_{meeting_id}_{timestamp}{ext}"
                    
                    # 确保目录存在
                    from django.conf import settings
                    upload_dir = os.path.join(settings.MEDIA_ROOT, 'meeting_recordings')
                    os.makedirs(upload_dir, exist_ok=True)
                    
                    # 保存文件
                    file_path = os.path.join(upload_dir, filename)
                    with open(file_path, 'wb+') as destination:
                        for chunk in audio_file.chunks():
                            destination.write(chunk)
                    
                    # 存储相对路径
                    audio_file_path = f"meeting_recordings/{filename}"
                elif meeting_record.attachments:
                    # 尝试从会议记录的附件中提取音频文件
                    import os
                    allowed_extensions = ['.mp3', '.wav', '.flac', '.ogg', '.aac']
                    for attachment in meeting_record.attachments.split(','):
                        if any(attachment.lower().endswith(ext) for ext in allowed_extensions):
                            audio_file_path = attachment
                            break
                
                if not audio_file_path:
                    return JsonResponse({'code': 1, 'msg': '请上传音频文件或选择已有音频的会议记录'})
                
                # 调用AI服务生成会议纪要
                try:
                    from apps.ai.utils.analysis_tools import default_meeting_analysis_tool
                    
                    # 构建会议信息
                    meeting_data = {
                        'title': meeting_record.title,
                        'meeting_date': meeting_record.meeting_date.strftime('%Y-%m-%d %H:%M:%S') if meeting_record.meeting_date else '',
                        'anchor_id': meeting_record.host.id if meeting_record.host else '',
                        'anchor_name': meeting_record.host.name if meeting_record.host else '',
                        'recorder_id': meeting_record.recorder.id if meeting_record.recorder else '',
                        'recorder_name': meeting_record.recorder.name if meeting_record.recorder else '',
                        'meeting_address': meeting_record.location if meeting_record.location else '',
                        'meeting_content': meeting_record.content,
                        'resolutions': meeting_record.resolution,
                        'join_uids': ','.join([str(user.id) for user in meeting_record.participants.all()]) if meeting_record.participants.exists() else '',
                        'join_names': ','.join([user.name for user in meeting_record.participants.all()]) if meeting_record.participants.exists() else '',
                        'sign_uids': ','.join([str(user.id) for user in meeting_record.attendees.all()]) if meeting_record.attendees.exists() else '',
                        'sign_names': ','.join([user.name for user in meeting_record.attendees.all()]) if meeting_record.attendees.exists() else '',
                        'share_uids': ','.join([str(user.id) for user in meeting_record.shared_users.all()]) if meeting_record.shared_users.exists() else '',
                        'share_names': ','.join([user.name for user in meeting_record.shared_users.all()]) if meeting_record.shared_users.exists() else '',
                        'audio_file_path': audio_file_path
                    }
                    
                    # 调用AI生成纪要
                    ai_result = default_meeting_analysis_tool.generate_meeting_minutes(meeting_data)
                    
                    return JsonResponse({'code': 0, 'data': ai_result})
                except Exception as e:
                    import logging
                    logging.error(f"AI生成会议纪要失败: {str(e)}")
                    return JsonResponse({'code': 1, 'msg': f'生成会议纪要失败：{str(e)}'})
            except Exception as e:
                import logging
                logger = logging.getLogger(__name__)
                logger.error(f"处理请求时发生异常：{str(e)}")
                return JsonResponse({'code': 1, 'msg': f'处理请求失败：{str(e)}'})

        # 普通保存请求
        form = MeetingMinutesForm(request.POST, instance=minute)
        if form.is_valid():
            try:
                minute = form.save(commit=False)
                minute.user = request.user
                minute.recorder = request.user
                
                # 如果有会议ID参数，关联会议记录
                meeting_id = request.POST.get('meeting_id') or request.GET.get('meeting_id')
                if meeting_id:
                    try:
                        from apps.oa.models import MeetingRecord
                        meeting_record = MeetingRecord.objects.get(pk=meeting_id)
                        minute.meeting_record = meeting_record
                    except MeetingRecord.DoesNotExist:
                        import logging
                        logger = logging.getLogger(__name__)
                        logger.warning(f"会议记录不存在：会议ID {meeting_id} 未找到对应的MeetingRecord")
                
                minute.save()
                messages.success(request, '会议纪要保存成功！')
                return redirect('personal:minutes_list')
            except Exception as e:
                messages.error(request, f'保存会议纪要时出错：{str(e)}')
        else:
            # 显示表单错误信息
            for field, errors in form.errors.items():
                for error in errors:
                    messages.error(request, f'{field}: {error}')
    else:
        # 如果是新建会议纪要，设置默认值
        if not minute:
            now = timezone.now()
            initial = {
                'meeting_date': now.strftime('%Y-%m-%dT%H:%M'),
                'meeting_type': 'regular',
                'is_public': True
            }
            
            # 检查是否有会议ID参数，有则填充表单
            meeting_id = request.GET.get('meeting_id')
            if meeting_id:
                try:
                    # 从会议记录获取信息
                    from apps.oa.models import MeetingRecord
                    meeting_record = MeetingRecord.objects.get(pk=meeting_id)
                    # 填充表单初始值
                    initial.update({
                        'title': meeting_record.title,
                        'meeting_date': meeting_record.meeting_date.strftime('%Y-%m-%dT%H:%M'),
                        'location': meeting_record.location if meeting_record.location else '',
                        'host': meeting_record.host.name if meeting_record.host else '',
                        'attendees': ', '.join([user.name for user in meeting_record.participants.all()]) if meeting_record.participants.exists() else ''
                    })
                except MeetingRecord.DoesNotExist:
                    messages.warning(request, '未找到对应的会议记录，使用默认值创建')
            # 检查是否有预约ID参数，有则填充表单
            elif request.GET.get('reservation_id'):
                reservation_id = request.GET.get('reservation_id')
                try:
                    # 从会议室预约获取信息
                    from apps.system.models import MeetingReservation
                    reservation = MeetingReservation.objects.get(pk=reservation_id)
                    # 安全地获取会议室名称，避免NoneType错误
                    meeting_room_name = reservation.meeting_room.name if reservation.meeting_room else ''
                    
                    # 尝试查找是否已有基于此预约创建的会议记录
                    from apps.oa.models import MeetingRecord
                    meeting_record = MeetingRecord.objects.filter(title=reservation.title, 
                                                                 meeting_date=reservation.start_time).first()
                    
                    # 如果没有找到关联的会议记录，自动创建一个
                    if not meeting_record:
                        meeting_record = MeetingRecord.objects.create(
                            title=reservation.title,
                            meeting_date=reservation.start_time,
                            meeting_end_time=reservation.end_time,
                            location=meeting_room_name,
                            host=reservation.organizer,
                            content=reservation.description or ''
                        )
                        # 添加参与者 - 确保正确处理ManyToMany关系
                        participants_list = list(reservation.attendees.all())
                        if participants_list:
                            meeting_record.participants.set(participants_list)
                        
                        # 记录日志
                        import logging
                        logger = logging.getLogger(__name__)
                        logger.info(f"基于预约ID {reservation_id} 自动创建了会议记录 {meeting_record.id}")
                        
                        # 将自动创建的会议记录ID存储在会话中，以便后续使用
                        request.session['auto_created_meeting_id'] = meeting_record.id
                    
                    try:
                        # 填充表单初始值
                        initial.update({
                            'title': reservation.title,
                            'meeting_date': reservation.start_time.strftime('%Y-%m-%dT%H:%M'),
                            'location': meeting_room_name,
                            'host': reservation.organizer.name if reservation.organizer else '',
                            'attendees': ', '.join([user.name for user in reservation.attendees.all()]) if reservation.attendees.exists() else '',
                            'meeting_id': meeting_record.id,  # 设置会议ID，确保AI生成功能可以使用
                            'content': f"本次会议使用{meeting_room_name}会议室，时间为{reservation.start_time.strftime('%Y-%m-%d %H:%M')}至{reservation.end_time.strftime('%H:%M')}"
                        })
                    except Exception as e:
                        # 记录具体错误信息以便调试
                        import logging
                        logging.error(f"填充会议室预约信息时出错: {str(e)}")
                        messages.warning(request, f'处理会议室预约信息时出现问题: {str(e)}')
                except MeetingReservation.DoesNotExist:
                    messages.warning(request, '未找到对应的会议室预约，使用默认值创建')
                except Exception as e:
                    # 更详细地捕获其他异常
                    import logging
                    logging.error(f"处理reservation_id时出错: {str(e)}")
                    messages.warning(request, f'处理会议室预约信息时出现错误: {str(e)}')
            
            form = MeetingMinutesForm(initial=initial)
        else:
            form = MeetingMinutesForm(instance=minute)
    
    # 获取可用的会议记录列表（只显示用户有权限的）
    from apps.oa.models import MeetingRecord
    user_meetings = MeetingRecord.objects.filter(
        Q(host=request.user) | Q(recorder=request.user) | Q(participants=request.user)
    ).order_by('-meeting_date')
    
    # 检查是否有预约ID参数并添加到上下文
    reservation_id = request.GET.get('reservation_id')
    context = {'form': form, 'minute': minute, 'user_meetings': user_meetings}
    if reservation_id:
        context['reservation_id'] = reservation_id
    
    # 检查是否有自动创建的会议ID，添加到上下文
    auto_meeting_id = request.session.get('auto_created_meeting_id')
    if auto_meeting_id:
        context['auto_meeting_id'] = auto_meeting_id
        # 移除会话中的临时存储
        del request.session['auto_created_meeting_id']
    return render(request, 'personal/minutes/form.html', context)


@login_required
def minutes_delete(request, pk):
    """删除会议纪要"""
    minute = get_object_or_404(MeetingMinutes, pk=pk)
    # 检查权限：只有创建者可以删除
    if minute.user != request.user:
        messages.error(request, '您没有权限删除此会议纪要！')
        return redirect('personal:minutes_list')
        
    if request.method == 'POST':
        minute.delete()
        messages.success(request, '会议纪要删除成功！')
        return redirect('personal:minutes_list')
    return render(request, 'personal/minutes/delete.html', {'minute': minute})


@login_required
def generate_minutes_word(request, pk):
    """
    生成会议纪要Word文档并下载
    根据固定格式要求，包含公司标志、会议信息表格和会议事项表格
    """
    minutes = get_object_or_404(MeetingMinutes, pk=pk)
    
    # 权限检查
    if minutes.user != request.user and not minutes.is_public:
        return HttpResponse('没有权限访问此会议纪要')
    
    # 创建Word文档
    doc = Document()
    
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    
    # 添加标题行（公司标志和名称）
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 尝试添加公司标志（如果存在）
    logo_path = os.path.join(settings.STATIC_ROOT, 'img', 'rdf.png')
    if os.path.exists(logo_path):
        try:
            # 先尝试在标题行中添加图片
            logo_run = header.add_run()
            logo_run.add_picture(logo_path, width=Inches(1.5))
        except:
            # 如果添加图片失败，继续执行
            pass
    
    # 添加公司名称
    company_name = header.add_run('江苏瑞德丰精密技术股份有限公司')
    company_name.font.name = '微软雅黑'
    company_name.font.size = Pt(16)
    company_name.bold = True
    
    # 添加副标题
    subtitle = doc.add_paragraph('会议纪要')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.name = '微软雅黑'
    subtitle_run.font.size = Pt(14)
    subtitle_run.bold = True
    
    # 添加基本信息表格（按照固定格式要求）
    info_table = doc.add_table(rows=4, cols=4)
    info_table.style = 'Table Grid'
    
    # 设置表格内容，严格按照固定格式要求
    info_table.cell(0, 0).text = '会议议题'
    info_table.cell(0, 1).text = minutes.title
    info_table.cell(0, 1).merge(info_table.cell(0, 2))  # 合并单元格使议题占两列
    info_table.cell(0, 3).text = '会议时间'
    
    # 格式化会议时间
    meeting_date_str = minutes.meeting_date.strftime('%Y年%m月%d日 %H:%M')
    info_table.cell(1, 3).text = meeting_date_str
    
    info_table.cell(1, 0).text = '会议类型'
    # 获取会议类型的中文名称
    meeting_type_dict = dict(MeetingTypeChoices.choices)
    info_table.cell(1, 1).text = meeting_type_dict.get(minutes.meeting_type, minutes.meeting_type)
    info_table.cell(1, 1).merge(info_table.cell(1, 2))  # 合并单元格使类型占两列
    
    info_table.cell(2, 0).text = '参会人员'
    info_table.cell(2, 1).text = minutes.attendees or ''
    info_table.cell(2, 1).merge(info_table.cell(2, 2))  # 合并单元格使参会人员占两列
    info_table.cell(2, 3).text = '会议主持'
    info_table.cell(3, 3).text = minutes.host or ''
    
    # 设置表格中第一列的样式为粗体
    for i in range(4):
        # 安全地设置粗体样式，避免索引错误
        for col in [0, 3]:
            if info_table.cell(i, col).paragraphs and info_table.cell(i, col).paragraphs[0].runs:
                info_table.cell(i, col).paragraphs[0].runs[0].bold = True
    
    # 添加会议记录人和审核人信息
    recorder_table = doc.add_table(rows=1, cols=4)
    recorder_table.style = 'Table Grid'
    recorder_table.cell(0, 0).text = '记录人：'
    recorder_table.cell(0, 1).text = minutes.recorder.name or minutes.recorder.username
    recorder_table.cell(0, 2).text = '审核人：'
    recorder_table.cell(0, 3).text = ''
    
    # 设置记录人表格中标签的样式为粗体
    for col in [0, 2]:
        if recorder_table.cell(0, col).paragraphs and recorder_table.cell(0, col).paragraphs[0].runs:
            recorder_table.cell(0, col).paragraphs[0].runs[0].bold = True
    
    # 添加空行
    doc.add_paragraph()
    
    # 添加会议事项表格（包含实际完成时间列）
    doc.add_paragraph('会议事项：')
    tasks_table = doc.add_table(rows=1, cols=5)  # 5列：No.、会议事项、责任人、计划完成时间、实际完成时间
    tasks_table.style = 'Table Grid'
    
    # 设置表头
    tasks_table.cell(0, 0).text = 'No.'
    tasks_table.cell(0, 1).text = '会议事项（措施）'
    tasks_table.cell(0, 2).text = '责任人'
    tasks_table.cell(0, 3).text = '计划完成时间'
    tasks_table.cell(0, 4).text = '实际完成时间'
    
    # 设置表头样式为粗体
    for cell in tasks_table.rows[0].cells:
        if cell.paragraphs and cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].bold = True
    
    # 处理会议决议和行动项
    # 优先使用decisions字段，如果为空则使用action_items
    items_source = minutes.decisions or getattr(minutes, 'action_items', None)
    
    if items_source:
        # 新的处理逻辑：解析会议决议格式，将标题、决策内容、目标合并为会议事项
        
        import re
        import logging
        logger = logging.getLogger(__name__)
        
        # 更加灵活的正则表达式，考虑各种可能的空格和换行变化
        # 使用re.DOTALL让.匹配换行符，并使用非贪婪匹配
        resolution_patterns = [
            # 主要模式：标准格式
            r'(\d+)\.\s*(.+?)\s*-\s*决策内容\s*：\s*(.+?)\s*-\s*执行对象\s*：\s*(.+?)\s*-\s*目标\s*：\s*(.+?)(?=\n\d+\.|\Z)',
            # 备用模式：处理格式可能的变化
            r'(\d+)\.\s*(.+?)\s*\n-\s*决策内容\s*：\s*(.+?)\s*\n-\s*执行对象\s*：\s*(.+?)\s*\n-\s*目标\s*：\s*(.+?)(?=\n\d+\.|\Z)'
        ]
        
        found_resolutions = False
        for pattern in resolution_patterns:
            # 使用re.DOTALL标志让.可以匹配换行符
            resolutions = re.findall(pattern, items_source, re.DOTALL)
            if resolutions:
                found_resolutions = True
                logger.info(f"使用正则表达式模式匹配到{len(resolutions)}个会议决议项")
                break
        
        if found_resolutions and resolutions:
            # 如果匹配到格式化的会议决议
            for res in resolutions:
                if len(res) >= 5:
                    _, title, decision_content, executor, target = res[:5]
                    
                    # 清理提取的内容，去除多余的空白字符
                    title = title.strip()
                    decision_content = decision_content.strip().rstrip('。')
                    executor = executor.strip()
                    target = target.strip().rstrip('。')
                    
                    # 合并标题、决策内容和目标到会议事项列
                    combined_item = f"{title}。{decision_content}。{target}"
                    
                    row_cells = tasks_table.add_row().cells
                    row_cells[0].text = str(len(tasks_table.rows) - 1)  # 动态计算序号
                    row_cells[1].text = combined_item
                    row_cells[2].text = executor  # 执行对象作为责任人
                    row_cells[3].text = ''
                    row_cells[4].text = ''
        else:
            # 旧格式兼容：按换行符分割文本
            logger.info("未匹配到格式化的会议决议，使用旧格式处理")
            items = items_source.split('\n')
            for i, item in enumerate(items, 1):
                if item.strip():
                    row_cells = tasks_table.add_row().cells
                    row_cells[0].text = str(i)
                    row_cells[1].text = item.strip()
                    row_cells[2].text = ''
                    row_cells[3].text = ''
                    row_cells[4].text = ''
    else:
        # 如果没有决议和行动项，添加一行空数据
        row_cells = tasks_table.add_row().cells
        row_cells[0].text = '1'
        row_cells[1].text = ''
        row_cells[2].text = ''
        row_cells[3].text = ''
        row_cells[4].text = ''
    
    # 设置列宽
    tasks_table.columns[0].width = Inches(0.5)
    tasks_table.columns[1].width = Inches(2.5)
    tasks_table.columns[2].width = Inches(1.0)
    tasks_table.columns[3].width = Inches(1.25)
    tasks_table.columns[4].width = Inches(1.25)
    
    # 创建响应
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{minutes.title}_会议纪要.docx"'
    
    # 保存文档到响应
    doc.save(response)
    
    return response

@login_required
def generate_minutes_preview(request, pk):
    """
    预览会议纪要Word文档
    使用网盘中的OfficePreviewHandler实现在线预览功能
    """
    minutes = get_object_or_404(MeetingMinutes, pk=pk)
    
    # 检查权限
    if minutes.user != request.user and not minutes.is_public:
        return HttpResponse('没有权限预览此会议纪要')
    
    # 生成临时Word文档
    temp_doc_path = None
    try:
        # 创建临时文件，确保使用正确的.docx扩展名
        fd, temp_doc_path = tempfile.mkstemp(suffix='.docx')
        os.close(fd)  # 关闭文件描述符，让python-docx来处理文件
        
        # 使用与下载功能相同的文档生成逻辑
        doc = Document()
        
        # 设置页面边距
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)
        
        # 添加标题行（公司标志和名称）
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 尝试添加公司标志（如果存在）
        logo_path = os.path.join(settings.STATIC_ROOT, 'img', 'rdf.png')
        if os.path.exists(logo_path):
            try:
                # 先尝试在标题行中添加图片
                logo_run = header.add_run()
                logo_run.add_picture(logo_path, width=Inches(1.5))
            except:
                # 如果添加图片失败，继续执行
                pass
        
        # 添加公司名称
        company_name = header.add_run('江苏瑞德丰精密技术股份有限公司')
        company_name.font.name = '微软雅黑'
        company_name.font.size = Pt(16)
        company_name.bold = True
        
        # 添加副标题
        subtitle = doc.add_paragraph('会议纪要')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.runs[0]
        subtitle_run.font.name = '微软雅黑'
        subtitle_run.font.size = Pt(14)
        subtitle_run.bold = True
        
        # 添加基本信息表格（按照固定格式要求）
        info_table = doc.add_table(rows=4, cols=4)
        info_table.style = 'Table Grid'
        
        # 设置表格内容，严格按照固定格式要求
        info_table.cell(0, 0).text = '会议议题'
        info_table.cell(0, 1).text = minutes.title
        info_table.cell(0, 1).merge(info_table.cell(0, 2))  # 合并单元格使议题占两列
        info_table.cell(0, 3).text = '会议时间'
        
        # 格式化会议时间
        meeting_date_str = minutes.meeting_date.strftime('%Y年%m月%d日 %H:%M')
        info_table.cell(1, 3).text = meeting_date_str
        
        info_table.cell(1, 0).text = '会议类型'
        # 获取会议类型的中文名称
        meeting_type_dict = dict(MeetingTypeChoices.choices)
        info_table.cell(1, 1).text = meeting_type_dict.get(minutes.meeting_type, minutes.meeting_type)
        info_table.cell(1, 1).merge(info_table.cell(1, 2))  # 合并单元格使类型占两列
        
        info_table.cell(2, 0).text = '参会人员'
        info_table.cell(2, 1).text = minutes.attendees or ''
        info_table.cell(2, 1).merge(info_table.cell(2, 2))  # 合并单元格使参会人员占两列
        info_table.cell(2, 3).text = '会议主持'
        info_table.cell(3, 3).text = minutes.host or ''
        
        # 设置表格中第一列的样式为粗体
        for i in range(4):
            # 安全地设置粗体样式，避免索引错误
            for col in [0, 3]:
                if info_table.cell(i, col).paragraphs and info_table.cell(i, col).paragraphs[0].runs:
                    info_table.cell(i, col).paragraphs[0].runs[0].bold = True
        
        # 添加会议记录人和审核人信息
        recorder_table = doc.add_table(rows=1, cols=4)
        recorder_table.style = 'Table Grid'
        recorder_table.cell(0, 0).text = '记录人：'
        recorder_table.cell(0, 1).text = minutes.recorder.name or minutes.recorder.username
        recorder_table.cell(0, 2).text = '审核人：'
        recorder_table.cell(0, 3).text = ''
        
        # 设置记录人表格中标签的样式为粗体
        for col in [0, 2]:
            if recorder_table.cell(0, col).paragraphs and recorder_table.cell(0, col).paragraphs[0].runs:
                recorder_table.cell(0, col).paragraphs[0].runs[0].bold = True
        
        # 添加空行
        doc.add_paragraph()
        
        # 添加会议事项表格（包含实际完成时间列）
        doc.add_paragraph('会议事项：')
        tasks_table = doc.add_table(rows=1, cols=5)  # 5列：No.、会议事项、责任人、计划完成时间、实际完成时间
        tasks_table.style = 'Table Grid'
        
        # 设置表头
        tasks_table.cell(0, 0).text = 'No.'
        tasks_table.cell(0, 1).text = '会议事项（措施）'
        tasks_table.cell(0, 2).text = '责任人'
        tasks_table.cell(0, 3).text = '计划完成时间'
        tasks_table.cell(0, 4).text = '实际完成时间'
        
        # 设置表头样式为粗体
        for cell in tasks_table.rows[0].cells:
            if cell.paragraphs and cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].bold = True
        
        # 处理会议决议和行动项
        import re
        import logging
        
        # 优先使用decisions字段，如果为空则使用action_items
        items_source = minutes.decisions or getattr(minutes, 'action_items', None)
        
        if items_source:
            # 定义多种正则表达式模式以增强匹配灵活性
            patterns = [
                # 主要模式：完整格式
                r'标题\s*[:：]\s*(.+?)\s*决策内容\s*[:：]\s*(.+?)\s*执行对象\s*[:：]\s*(.+?)\s*目标\s*[:：]\s*(.+)',
                # 备用模式1：可能有缺失字段
                r'标题\s*[:：]\s*(.+?)\s*决策内容\s*[:：]\s*(.+?)\s*(?:执行对象\s*[:：]\s*(.+?))?\s*(?:目标\s*[:：]\s*(.+?))?',
                # 备用模式2：简化格式
                r'标题\s*[:：]\s*(.+?)\s*(决策内容|内容)\s*[:：]\s*(.+?)(?:\n|$)',
            ]
            
            # 初始化计数器
            item_count = 0
            
            # 尝试使用正则表达式解析结构化数据
            match_found = False
            
            for pattern in patterns:
                matches = re.finditer(pattern, items_source, re.DOTALL)
                for match in matches:
                    match_found = True
                    item_count += 1
                    
                    # 提取匹配的内容组
                    groups = match.groups()
                    
                    # 初始化变量
                    title = ''
                    content = ''
                    executor = ''
                    target = ''
                    
                    # 根据匹配的组数分配值
                    if len(groups) >= 1:
                        title = groups[0].strip()
                    if len(groups) >= 2:
                        # 如果第二组是'决策内容'或'内容'标签，则实际内容在第三组
                        if groups[1] in ['决策内容', '内容'] and len(groups) >= 3:
                            content = groups[2].strip()
                            if len(groups) >= 4:
                                executor = groups[3].strip()
                            if len(groups) >= 5:
                                target = groups[4].strip()
                        else:
                            content = groups[1].strip()
                            if len(groups) >= 3:
                                executor = groups[2].strip()
                            if len(groups) >= 4:
                                target = groups[3].strip()
                    
                    # 合并标题、决策内容和目标为会议事项
                    combined_item = title
                    if content:
                        combined_item += ' ' + content
                    if target:
                        combined_item += ' ' + target
                    
                    # 清理多余的空白字符
                    combined_item = ' '.join(combined_item.split())
                    
                    # 添加到表格
                    row_cells = tasks_table.add_row().cells
                    row_cells[0].text = str(item_count)
                    row_cells[1].text = combined_item
                    row_cells[2].text = executor
                    row_cells[3].text = ''
                    row_cells[4].text = ''
                    
                    logging.info(f"预览生成 - 解析到会议决议项: {combined_item}, 责任人: {executor}")
                
                # 如果找到匹配项，不再尝试其他模式
                if match_found:
                    break
            
            # 如果没有找到结构化数据，回退到按行分割的旧方式
            if not match_found:
                logging.info("预览生成 - 未找到结构化会议决议，使用按行分割方式")
                items = items_source.split('\n')
                for i, item in enumerate(items, 1):
                    if item.strip():
                        row_cells = tasks_table.add_row().cells
                        row_cells[0].text = str(i)
                        row_cells[1].text = item.strip()
                        row_cells[2].text = ''
                        row_cells[3].text = ''
                        row_cells[4].text = ''
        else:
            # 如果没有决议和行动项，添加一行空数据
            row_cells = tasks_table.add_row().cells
            row_cells[0].text = '1'
            row_cells[1].text = ''
            row_cells[2].text = ''
            row_cells[3].text = ''
            row_cells[4].text = ''
        
        # 设置列宽
        tasks_table.columns[0].width = Inches(0.5)
        tasks_table.columns[1].width = Inches(2.5)
        tasks_table.columns[2].width = Inches(1.0)
        tasks_table.columns[3].width = Inches(1.25)
        tasks_table.columns[4].width = Inches(1.25)
        
        # 保存文档到临时文件
        doc.save(temp_doc_path)
        
        # 使用网盘中的OfficePreviewHandler处理预览
        # 不指定conversion_format，让处理器自动选择合适的预览方式
        preview_result = OfficePreviewHandler.preview_office_file(temp_doc_path)
        
        # 添加文件ID和下载URL信息，与disk应用保持一致
        preview_result['file_id'] = pk
        preview_result['download_url'] = request.build_absolute_uri(
            reverse('personal:minutes_download', kwargs={'pk': pk})
        )
        
        # 构建预览页面的响应
        context = {
            'preview_data': preview_result,
            'filename': f"{minutes.title}_会议纪要预览.docx",
            'minute': minutes
        }
        
        # 返回渲染的预览页面
        return render(request, 'personal/minutes/preview.html', context)
    except Exception as e:
        # 处理异常情况
        import logging
        logger = logging.getLogger('personal')
        logger.error(f'会议纪要预览失败: {str(e)}', exc_info=True)
        return HttpResponse(f'预览出错: {str(e)}')
    finally:
        # 清理临时文件
        try:
            if temp_doc_path and os.path.exists(temp_doc_path):
                os.remove(temp_doc_path)
        except Exception as e:
            import logging
            logger = logging.getLogger('personal')
            logger.warning(f'清理临时文件失败: {str(e)}')

@login_required
def minutes_detail(request, pk):
    """会议纪要详情"""
    minute = get_object_or_404(MeetingMinutes, pk=pk)
    # 检查权限：非公开的会议纪要只有创建者可以查看
    if not minute.is_public and minute.user != request.user:
        messages.error(request, '您没有权限查看此会议纪要！')
        return redirect('personal:minutes_list')
        
    context = {'minute': minute}
    return render(request, 'personal/minutes/detail.html', context)